"""
update_docx.py — Rebuilds 4 DOCX files from their corresponding HTML source files.
Uses BeautifulSoup to parse HTML and python-docx to write structured DOCX files.
"""

import re
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/alghali/Downloads/Compass"

# ─────────────────────────────────────────────────────────────────────────────
# Classes to skip entirely (decorative / chart / navigation only)
# ─────────────────────────────────────────────────────────────────────────────
SKIP_CLASS_SUBSTRINGS = [
    # Charts and diagrams
    'chart', 'canvas', 'tf-diagram', 'tf-timeline', 'tf-scenario',
    'tf-cycle', 'tf-col', 'tf-box', 'tf-day', 'tf-score', 'tf-new',
    'tf-absent', 'tf-obsolete', 'signal-compare', 'signal-vs',
    'decay-gauge', 'trajectory', 'signal-box',
    # Stats / decorative numbers
    'stat-row', 'stat-box', 'stat-strip', 'stat-card', 'stat-value', 'stat-label',
    # Icon / number decorations
    'card-icon', 'card-tags', 'auth-icon', 'step-num', 'use-num',
    'signal-num', 'phase-dot', 'phase-num', 'use-icon',
    'diff-icon', 'diff-num', 'product-icon', 'feature-icon',
    'zone-icon', 'section-icon', 'item-icon', 'item-num',
    'cat-number', 'cat-num', 'cat-icon',
    'stat-num', 'stat-chip',
    'audience-icon',
    # Badges and labels
    'badge', 'hero-badge', 'zone-badge', 'use-badge', 'we-badge',
    'zone-label', 'ev-strength',
    # Navigation
    'nav-logo', 'nav-section',
    # Score range labels (e.g. "76–100" pill next to zone card heading)
    'score-range', 'zone-range', 'zone-score',
    # Other decorative
    'we-result-score', 'we-label', 'we-result', 'we-row', 'we-grid',
    'decay-item', 'decay-hl', 'decay-zone', 'decay-sub',
    'sc-matrix', 'sc-matrix-cell', 'sc-matrix-sub', 'sc-zone',
    'alert-icon',
]

SKIP_CLASSES_EXACT = {
    'hero-badge', 'hero-meta', 'hero-meta-item',
}

# Inline CSS styles that indicate a purely decorative element
# (e.g. emoji icon divs with font-size only)
DECORATIVE_INLINE_STYLE_PATTERNS = [
    r'font-size:\s*\d+px.*margin-bottom',  # icon sizing
    r'justify-content:space-between.*align-items',  # flex row decorators
]

# Classes for alert/callout containers.
# These are rendered as indented italic paragraphs, with decorative children stripped.
ALERT_CONTAINER_CLASSES = {
    'alert', 'alert-box', 'callout-box', 'note', 'info-box',
    'callout', 'warning', 'key-insight', 'clinical-note',
    'quote-block', 'insight', 'cite-block', 'reference',
}


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def clean_text(text):
    """Normalise whitespace."""
    return re.sub(r'\s+', ' ', text or '').strip()


def heading_text(el):
    """
    Extract heading text from a heading element.
    Handles inline <br> by replacing with a space, then normalises whitespace.
    """
    # Replace <br> with space before extracting text
    for br in el.find_all('br'):
        br.replace_with(' ')
    return clean_text(el.get_text())


def is_decorative(el):
    """Return True if this element is purely decorative and should be skipped."""
    if not isinstance(el, Tag):
        return True
    classes = el.get('class', [])
    classes_str = ' '.join(classes)

    if any(sub in classes_str for sub in SKIP_CLASS_SUBSTRINGS):
        return True
    if set(classes) & SKIP_CLASSES_EXACT:
        return True

    # Skip divs that have inline styles indicating pure decoration:
    # e.g. <div style="font-size:22px; margin-bottom:10px;">🛡️</div>
    # or <div style="display:flex; justify-content:space-between; align-items:...">
    inline_style = el.get('style', '')
    if inline_style and el.name in ('div', 'span'):
        text = el.get_text(strip=True)
        # Single emoji or purely visual short text with font-size/margin style = icon
        if re.search(r'font-size:\s*\d+px', inline_style) and len(text) <= 5:
            return True
        # Flex row with space-between and no meaningful content other than badges
        if ('justify-content:space-between' in inline_style or
                'justify-content: space-between' in inline_style):
            # Skip if all children are decorative
            real_children = [c for c in el.children
                             if isinstance(c, Tag) and c.name not in ('br',)
                             and not is_decorative(c)]
            if not real_children:
                return True
        # Score range pill: small text, border-radius, background colour, no class
        if (not el.get('class') and
                re.search(r'border-radius', inline_style) and
                re.search(r'background', inline_style) and
                len(text) <= 12):
            return True
        # Decorative zone label spans: font-size ≤ 10px with letter-spacing/color,
        # short text (e.g. "DANGER", "DECEPTIVE")
        if (re.search(r'font-size:\s*(8|9|10|11)px', inline_style) and
                re.search(r'font-weight:\s*[6-9]00', inline_style) and
                len(text) <= 15):
            return True

    return False


def set_cell_bg(cell, hex_color):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_section_label(doc, label_text):
    """Add a small-caps section label paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(label_text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    return p


def process_inline(para, element):
    """
    Walk element children and add runs, preserving bold/italic/code/link inline
    formatting. Skips decorative child elements.
    """
    if isinstance(element, NavigableString):
        text = re.sub(r'\s+', ' ', str(element))
        if text.strip():
            para.add_run(text)
        return

    for child in element.children:
        if isinstance(child, NavigableString):
            text = re.sub(r'\s+', ' ', str(child))
            if text.strip():
                para.add_run(text)
        elif isinstance(child, Tag):
            # Skip decorative inline children
            if is_decorative(child):
                continue
            tag = child.name
            if tag in ('strong', 'b'):
                text = clean_text(child.get_text())
                if text:
                    r = para.add_run(text)
                    r.bold = True
            elif tag in ('em', 'i'):
                text = clean_text(child.get_text())
                if text:
                    r = para.add_run(text)
                    r.italic = True
            elif tag == 'code':
                text = clean_text(child.get_text())
                if text:
                    r = para.add_run(text)
                    r.font.name = 'Courier New'
                    r.font.size = Pt(9.5)
            elif tag == 'a':
                text = clean_text(child.get_text())
                if text:
                    para.add_run(text)
            elif tag == 'br':
                para.add_run(' ')  # treat <br> as space in inline context
            elif tag in ('span', 'sup', 'sub', 'small', 'mark', 'abbr', 'label'):
                process_inline(para, child)
            else:
                # Block-level in inline context — just get text
                text = clean_text(child.get_text())
                if text:
                    para.add_run(' ' + text)


def add_table_from_bs(doc, bs_table):
    """Convert a BeautifulSoup <table> element to a python-docx table."""
    rows = bs_table.find_all('tr')
    if not rows:
        return

    max_cols = 0
    for row in rows:
        cells = row.find_all(['th', 'td'])
        max_cols = max(max_cols, len(cells))

    if max_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    for r_idx, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        is_header_row = r_idx == 0 or all(c.name == 'th' for c in cells)

        for c_idx, cell in enumerate(cells):
            if c_idx >= max_cols:
                break
            cell_text = clean_text(cell.get_text())
            tc = table.rows[r_idx].cells[c_idx]
            para = tc.paragraphs[0]
            run = para.add_run(cell_text)
            if is_header_row:
                run.bold = True
                set_cell_bg(tc, 'D9E1F2')

    doc.add_paragraph()  # spacing after table


# ─────────────────────────────────────────────────────────────────────────────
# Main recursive converter
# ─────────────────────────────────────────────────────────────────────────────

def process_element(el, doc):
    """
    Recursively process a BS4 element and emit content into `doc`.
    """
    if not isinstance(el, Tag):
        return
    if is_decorative(el):
        return

    tag = el.name
    classes = set(el.get('class', []))

    # ── Skip footer/canvas entirely ───────────────────────────────────────────
    if tag in ('footer', 'canvas'):
        return

    # ── Headings ──────────────────────────────────────────────────────────────
    if tag == 'h1':
        text = heading_text(el)
        if text:
            doc.add_heading(text, level=1)
        return

    if tag == 'h2':
        text = heading_text(el)
        if text:
            doc.add_heading(text, level=2)
        return

    if tag == 'h3':
        text = heading_text(el)
        if text:
            doc.add_heading(text, level=3)
        return

    if tag == 'h4':
        text = heading_text(el)
        if text:
            doc.add_heading(text, level=4)
        return

    if tag in ('h5', 'h6'):
        text = heading_text(el)
        if text:
            p = doc.add_paragraph(style='Normal')
            r = p.add_run(text)
            r.bold = True
        return

    # ── Section label divs ────────────────────────────────────────────────────
    if 'section-label' in classes:
        text = clean_text(el.get_text())
        if text:
            add_section_label(doc, text)
        return

    # ── Hero container: only recurse into meaningful children ─────────────────
    if 'hero' in classes:
        for child in el.find_all(['h1', 'h2', 'h3', 'p'], recursive=True):
            process_element(child, doc)
        return

    # ── Alert / callout containers ────────────────────────────────────────────
    if tag in ('div', 'aside', 'section', 'article') and classes & ALERT_CONTAINER_CLASSES:
        # Recurse into child elements, but only render p/h3/h4 children
        # (skip decorative divs that contain only visual widgets)
        for child in el.children:
            if not isinstance(child, Tag):
                continue
            if is_decorative(child):
                continue
            ctag = child.name
            cclasses = set(child.get('class', []))
            if ctag == 'p':
                text = clean_text(child.get_text())
                if text:
                    p = doc.add_paragraph(style='Normal')
                    p.paragraph_format.left_indent = Inches(0.35)
                    process_inline(p, child)
            elif ctag in ('h3', 'h4', 'h5'):
                text = heading_text(child)
                if text:
                    lvl = {'h3': 3, 'h4': 4, 'h5': 4}.get(ctag, 4)
                    doc.add_heading(text, level=lvl)
            elif ctag in ('ul', 'ol'):
                process_element(child, doc)
            # Skip pure decorative div children (e.g., visual grid matrix)
        return

    # ── Paragraphs ────────────────────────────────────────────────────────────
    if tag == 'p':
        text = clean_text(el.get_text())
        if not text:
            return
        p = doc.add_paragraph(style='Normal')
        process_inline(p, el)
        return

    # ── Blockquote ────────────────────────────────────────────────────────────
    if tag == 'blockquote':
        text = clean_text(el.get_text())
        if text:
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            r = p.add_run(f'"{text}"')
            r.italic = True
        return

    # ── Unordered / ordered lists ─────────────────────────────────────────────
    if tag in ('ul', 'ol'):
        list_style = 'List Bullet' if tag == 'ul' else 'List Number'
        for li in el.find_all('li', recursive=False):
            # Each li might contain nested block elements (h4 + p pattern)
            # Pull out sub-headings and paragraphs inside
            sub_headings = li.find_all(['h3', 'h4', 'h5'], recursive=True)
            sub_paras = li.find_all('p', recursive=True)
            sub_lists = li.find_all(['ul', 'ol'], recursive=False)

            if sub_headings or sub_paras:
                # Render as heading/paragraph block, not a flat list item
                for sub in li.find_all(['h3', 'h4', 'h5', 'p', 'ul', 'ol'], recursive=False):
                    process_element(sub, doc)
                # If nothing found at top level, go one level deeper
                if not list(li.find_all(['h3', 'h4', 'h5', 'p', 'ul', 'ol'], recursive=False)):
                    for sub in li.children:
                        if isinstance(sub, Tag):
                            process_element(sub, doc)
            else:
                text = clean_text(li.get_text())
                if text:
                    p = doc.add_paragraph(style=list_style)
                    process_inline(p, li)
            # Nested lists
            for nested in sub_lists:
                for nested_li in nested.find_all('li', recursive=False):
                    nested_text = clean_text(nested_li.get_text())
                    if nested_text:
                        p = doc.add_paragraph(style='List Bullet 2' if tag == 'ul' else 'List Number 2')
                        process_inline(p, nested_li)
        return

    # ── Tables ────────────────────────────────────────────────────────────────
    if tag == 'table':
        add_table_from_bs(doc, el)
        return

    # ── Horizontal rule ───────────────────────────────────────────────────────
    if tag == 'hr':
        doc.add_paragraph()
        return

    # ── Generic containers: recurse ───────────────────────────────────────────
    if tag in ('div', 'section', 'article', 'main', 'aside', 'header',
               'figure', 'figcaption', 'details', 'summary', 'body',
               'form', 'fieldset'):

        # Check if this div has only text-level content (no block children)
        block_tags = {'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                      'ul', 'ol', 'table', 'div', 'section', 'article',
                      'blockquote', 'pre', 'figure'}
        block_children = [c for c in el.children
                          if isinstance(c, Tag) and c.name in block_tags
                          and not is_decorative(c)]

        if not block_children:
            # Treat as a single paragraph if it has text
            text = clean_text(el.get_text())
            if text and tag not in ('header',):
                p = doc.add_paragraph(style='Normal')
                process_inline(p, el)
            return

        # Recurse into children
        for child in el.children:
            if isinstance(child, Tag):
                process_element(child, doc)
        return


# ─────────────────────────────────────────────────────────────────────────────
# Per-file builder
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(html_path, out_path):
    print(f"\n{'='*60}")
    print(f"Processing: {html_path}")
    print(f"Output:     {out_path}")

    with open(html_path, encoding='utf-8') as f:
        html = f.read()

    soup = BeautifulSoup(html, 'html.parser')

    # Remove nav, script, style, footer upfront
    for tag in soup.find_all(['nav', 'script', 'style', 'footer', 'canvas']):
        tag.decompose()

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Default paragraph font ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Heading styles ──
    heading_specs = [(20, True), (16, True), (13, True), (12, True)]
    for level, (pt_size, bold) in enumerate(heading_specs, start=1):
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = 'Calibri'
        hstyle.font.size = Pt(pt_size)
        hstyle.font.bold = bold
        hstyle.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        hstyle.paragraph_format.space_before = Pt(16 if level == 1 else 10)
        hstyle.paragraph_format.space_after = Pt(4)

    # ── Process content ──
    main = soup.find('main') or soup.find('body')
    for child in main.children:
        if isinstance(child, Tag):
            process_element(child, doc)

    doc.save(out_path)
    print(f"  Saved — {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

MAPPINGS = [
    ('compass_framework.html',    'compass_framework.docx'),
    ('market_landscape.html',     'market_landscape.docx'),
    ('platform_description.html', 'NUR_Platform_Description.docx'),
    ('system_design.html',        'system_design.docx'),
]

if __name__ == '__main__':
    for html_name, docx_name in MAPPINGS:
        html_path = f"{BASE}/{html_name}"
        docx_path = f"{BASE}/{docx_name}"
        try:
            build_docx(html_path, docx_path)
        except Exception as e:
            import traceback
            print(f"\nERROR processing {html_name}:")
            traceback.print_exc()

    print("\nAll done.")
